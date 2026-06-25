# -*- coding: utf-8 -*-
"""Generate team division document matching the format of
项目组成员及任务分工-项目名称.docx, with 6-member content from the design spec/plan.

Usage: python generate_team_division.py
Output: 项目组成员及任务分工-高精度人脸检测算法及实践.docx
"""
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import lxml.etree as et

doc = Document()

for sec in doc.sections:
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.18)
    sec.right_margin = Cm(3.18)


def set_cn_font(run, cn_name='宋体_GB2312', size_emu=None, bold=None):
    rPr = run._element.get_or_add_rPr()
    rF = rPr.find(qn('w:rFonts'))
    if rF is None:
        rF = et.SubElement(rPr, qn('w:rFonts'))
    rF.set(qn('w:eastAsia'), cn_name)
    rF.set(qn('w:ascii'), 'Times New Roman')
    rF.set(qn('w:hAnsi'), 'Times New Roman')
    if size_emu is not None:
        run.font.size = size_emu
    if bold is not None:
        run.font.bold = bold


p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('      项目组成员及任务分工')
set_cn_font(r, '宋体_GB2312', Emu(203200), bold=True)

doc.add_paragraph()

# 6 members (A-F modules per plan) + 1 empty row = 8 rows (1 header + 6 member + 1 empty)
table = doc.add_table(rows=8, cols=4, style='Table Grid')
table.alignment = 1

headers = ['姓名', '项目报告中的主要贡献', '项目代码中的主要贡献', '成绩']
for ci, h_text in enumerate(headers):
    cell = table.rows[0].cells[ci]
    cell.text = ''
    run = cell.paragraphs[0].add_run(h_text)
    set_cn_font(run, '宋体_GB2312')

# Six members matching the plan's six modules (A-F)
members = [
    # ── A: 数据采集与标注 (Task 1-3) ──
    (
        '孙悟空',

        '1. 撰写第1章需求分析中数据采集部分的应用场景与前提条件分析\n'
        '2. 撰写第2章功能设计中数据采集与预处理模块的详细设计方案\n'
        '3. 完成数据标注格式设计文档（WIDER Face原始格式到YOLO格式转换规则）\n'
        '4. 撰写第2章数据设计中数据集来源与结构的分析（训练集12,880张/验证集3,226张）\n'
        '5. 编写可行性分析中的数据合规性论证（CC BY 4.0协议/WIDER Face使用许可）\n'
        '6. 编写第4章数据加载模块测试用例设计与测试结果记录',

        '1. data/download_widerface.py\n'
        '   - 实现 WIDER Face 数据集自动下载（Hugging Face/官方源）\n'
        '   - 实现 ZIP 自动解压与幂等下载（已存在则跳过）\n'
        '2. data/generate_list.py\n'
        '   - 解析 WIDER Face 原始标注文件（变长记录格式）\n'
        '   - 实现 (x,y,w,h) 到 YOLO (cx,cy,nw,nh) 归一化坐标转换\n'
        '   - 生成 train_list.txt / val_list.txt 图片列表文件\n'
        '3. data/check_annotation.py\n'
        '   - 实现四项自动质检规则（越界/尺寸<5px/面积>80%/面积<0.1%）\n'
        '   - 输出 quality_report.txt 质检报告\n'
        '4. 生成并维护 data/annotations/widerface.yaml 数据集配置文件',

        ''
    ),
    # ── B: 数据增强与加载 (Task 4-6) ──
    (
        '猪八戒',

        '1. 撰写第2章数据增强模块的详细设计方案（初期/后期两阶段增强策略）\n'
        '2. 撰写第3章数据增强模块的实现文档（Albumentations pipeline构建方法）\n'
        '3. 完成 DataLoader 模块的接口设计与数据流描述\n'
        '4. 撰写技术可行性分析中 Albumentations 技术栈的论证\n'
        '5. 编写第4章增强可视化测试的测试用例与验证标准\n'
        '6. 撰写第1章非功能需求中可扩展性需求（增强可配置性）的分析',

        '1. dataset/augmentation.py\n'
        '   - 实现 get_train_augmentation(phase) 工厂函数\n'
        '   - 初期增强策略（epoch1-50）：HorizontalFlip+BrightnessContrast+HueSaturationValue+Blur\n'
        '   - 后期增强策略（epoch51-100）：初期策略+CoarseDropout(Cutout遮挡)+强模糊\n'
        '   - 验证集增强策略：仅 Resize 640x640，无数据扰动\n'
        '   - 配置 BboxParams(format="yolo") 确保边界框坐标同步变换\n'
        '2. dataset/dataloader.py\n'
        '   - 实现 WiderFaceDataset 类（继承 torch.utils.data.Dataset）\n'
        '   - 实现 create_dataloader() 工厂函数（自动选择增强策略）\n'
        '   - 实现自定义 collate_fn 处理变长目标列表\n'
        '3. dataset/vis_aug.py\n'
        '   - 实现随机 batch 增强效果可视化，保存为 JPG 供人工确认',

        ''
    ),
    # ── C: 模型结构设计 (Task 7-9) ──
    (
        '张无忌',

        '1. 撰写第2章模型结构设计模块的详细设计文档\n'
        '2. 撰写 YOLOv8n 模型选型的技术论证（3.0M参数/CSPDarknet/PAN-FPN/Decoupled Head）\n'
        '3. 编写 model_config.yaml 配置文件中各参数的含义与设计依据\n'
        '4. 撰写第3章技术选型表（12项技术/框架的版本与选型理由）\n'
        '5. 编写技术可行性分析中深度学习框架部分的论证\n'
        '6. 绘制第2章系统总体架构设计中的四层架构图\n'
        '7. 编写第4章模型输出形状测试的测试用例',

        '1. model/model_config.yaml\n'
        '   - 配置模型选型（yolov8n-face.pt，ImageNet预训练权重）\n'
        '   - 配置输入尺寸（640x640）和网络结构（backbone/neck/head）\n'
        '   - 配置损失函数权重（box_loss=7.5/cls_loss=0.5/dfl_loss=1.5）\n'
        '   - 配置训练超参数（epochs/batch_size/optimizer/lr/warmup/amp）\n'
        '   - 配置增强控制参数（mosaic/mixup/hsv/fliplr/translate/scale）\n'
        '2. model/model_test.py\n'
        '   - 实现 test_model_output_shapes()：验证前向传播输出shape正确\n'
        '   - 实现 test_model_export_readiness()：验证TorchScript可导出\n'
        '   - 打印模型参数量（3,011,043 params）和计算量（8.2 GFLOPs）',

        ''
    ),
    # ── D: 训练与调优 (Task 8,10,11) ──
    (
        '令狐冲',

        '1. 撰写第2章模型训练模块的详细设计方案（两阶段训练策略）\n'
        '2. 撰写学习率调度策略的理论分析（Warmup+余弦退火至1e-5）\n'
        '3. 撰写混合精度训练（AMP）的技术原理与显存优化分析\n'
        '4. 编写第3章训练模块的实现文档（train.py/config.py核心逻辑）\n'
        '5. 撰写第1章非功能需求中性能需求（训练/推理性能指标）的分析\n'
        '6. 编写模型训练用例描述（UC02用例表）\n'
        '7. 撰写系统集成测试步骤2（模型训练阶段）的测试方案',

        '1. training/train.py\n'
        '   - 实现 train_v1()：YOLOv8n基线训练（100 epoch, AdamW, lr=1e-3, AMP, Mosaic）\n'
        '   - 实现 train_v2()：加载v1权重闭环微调（+30 epoch, lr=1e-4, Mosaic=0.5）\n'
        '   - 实现 prepare_data_yaml()：自动生成ultralytics格式数据配置文件\n'
        '   - 添加 cv2.imread Unicode路径兼容补丁（Windows中文路径适配）\n'
        '2. training/config.py\n'
        '   - 实现 load_config()：从model_config.yaml加载全部超参数\n'
        '   - 实现 print_training_summary()：训练前打印完整配置摘要\n'
        '   - 实现 compute_warmup_steps()：自动计算预热步数（>=500步保障）\n'
        '3. training/resume_train.py\n'
        '   - 实现从last.pt断点续训（resume=True）\n'
        '   - 支持训练中断后无缝恢复（保留优化器状态和epoch计数）\n'
        '4. 集成 TensorBoard 训练监控：实时记录 Loss/LR/mAP 曲线',

        ''
    ),
    # ── E: 评估与错误分析 (Task 12-13) ──
    (
        '欧阳峰',

        '1. 撰写第2章模型评估模块的详细设计方案（多维度评估指标体系）\n'
        '2. 撰写困难样本分析算法设计（IoU贪心匹配+人脸尺度四级分类+错误分布统计）\n'
        '3. 编写 v1 vs v2 闭环提升对比分析的方法论描述\n'
        '4. 撰写第3章评估模块的实现文档（evaluate.py/analyze_errors.py）\n'
        '5. 撰写可行性分析中评估工具的技术论证（widerface-evaluate/scikit-learn）\n'
        '6. 编写模型评估用例描述（UC03用例表）\n'
        '7. 编写第4章评估与性能测试用例（ONNX精度验证+基准测试+错误分析测试）\n'
        '8. 撰写系统集成测试结论，汇总全流程测试结果与局限性分析',

        '1. evaluation/evaluate.py\n'
        '   - 实现 evaluate_widerface()：调用model.val()计算mAP50/mAP50-95/Precision/Recall\n'
        '   - 实现 plot_pr_curve()：生成并保存P-R曲线图\n'
        '   - 实现 save_metrics()：评估指标导出为JSON文件（metrics_v1.json/metrics_v2.json）\n'
        '   - 实现 v1 vs v2 指标对比表（控制台格式化输出四列对比）\n'
        '2. evaluation/analyze_errors.py\n'
        '   - 实现 analyze_errors()：逐图推理+IoU=0.5贪心匹配算法\n'
        '   - 识别漏检（False Negative）和误检（False Positive）\n'
        '   - 按人脸尺度四级分类统计（Tiny<32px/Small/Medium/Large>256px）\n'
        '   - 实现 plot_error_visualization()：TOP20困难样本拼接图（FN红色/FP蓝色标注）\n'
        '   - 实现 plot_error_distribution()：错误分布柱状图\n'
        '   - 输出 false_negatives_top20.txt / false_positives_top20.txt 文本报告',

        ''
    ),
    # ── F: 部署与集成 (Task 14-18) ──
    (
        '韦小宝',

        '1. 撰写第2章部署模块的详细设计方案（ONNX导出流程+实时检测架构）\n'
        '2. 撰写 ONNX 模型导出与精度验证的技术方案（FP32/FP16/onnx-simplifier）\n'
        '3. 撰写模型部署的跨平台可移植性分析（CPU/GPU/移动端三种部署模式）\n'
        '4. 编写第3章部署模块的实现文档（export_onnx.py/benchmark.py/realtime_detect.py）\n'
        '5. 撰写第1章非功能需求中可移植性需求的分析（ONNX Runtime多后端支持）\n'
        '6. 编写实时检测用例描述（UC04用例表）\n'
        '7. 编写第4章部署阶段测试用例（ONNX导出测试+性能基准测试+实时检测集成测试）\n'
        '8. 完成 README.md 全项目说明文档（安装/使用/性能指标/技术栈）',

        '1. deployment/export_onnx.py\n'
        '   - 实现 export_to_onnx()：调用model.export()导出ONNX模型\n'
        '   - 支持 FP32 和 FP16 两种精度导出（FP16模型大小6.17MB <=10MB）\n'
        '   - 实现 validate_precision()：100张测试图PyTorch vs ONNX MAE精度验证\n'
        '   - 使用 onnx-simplifier 简化模型计算图，减少冗余运算\n'
        '2. deployment/benchmark.py\n'
        '   - 实现 benchmark_pytorch()：PyTorch GPU推理速度测试（50次预热+200次测量）\n'
        '   - 实现 benchmark_onnx()：ONNX Runtime多后端推理速度测试\n'
        '   - 实现 get_model_size()：模型大小检查（<=10MB达标验证）\n'
        '   - 输出 FPS >= 30、Size <= 10MB 达标判定\n'
        '3. deployment/realtime_detect.py\n'
        '   - 实现 run_realtime()：OpenCV读取摄像头/视频流+逐帧YOLO推理\n'
        '   - 实时绘制绿色人脸框+置信度分数+左上角FPS/人脸数叠加显示\n'
        '   - 使用deque滑动窗口计算30帧平均FPS，消除瞬时波动\n'
        '   - 按q键优雅退出并释放摄像头资源\n'
        '4. README.md\n'
        '   - 编写完整项目说明文档（快速开始/项目结构/模型性能/使用指南/技术栈）\n'
        '   - 将所有 TBD 占位符替换为实际训练和测试数据',

        ''
    ),
    # ── Empty row (保持与原模板格式一致) ──
    (
        '',
        '',
        '',
        ''
    ),
]

for ri, (name, report_contrib, code_contrib, grade) in enumerate(members):
    row = table.rows[ri + 1]
    for ci, val in enumerate([name, report_contrib, code_contrib, grade]):
        cell = row.cells[ci]
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        set_cn_font(run, '宋体_GB2312')

output_path = 'C:/学习/课程设计/项目组成员及任务分工-高精度人脸检测算法及实践.docx'
doc.save(output_path)

import os
print(f'OK -> {output_path}')
print(f'Size: {os.path.getsize(output_path)/1024:.1f} KB')
print('Done. 6 members (A-F) matching the design plan.')
